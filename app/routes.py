"""
GenoPilot – Rutas Flask
- '/'         : formulario
- '/generate' : genera informe (intenta PDF vía DOCX→PDF; si no, entrega DOCX)
Incluye fix de "tabla que desborda": tras renderizar el DOCX se fuerza el ancho
de la tabla de resultados al ancho útil de página, con layout fijo y soft-wrap.
"""
from __future__ import annotations

from flask import Blueprint, render_template, request, send_file
import os, json, datetime, re, tempfile, unicodedata
from io import BytesIO

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import os, platform
DISABLE_PDF = os.getenv("DISABLE_PDF", "0") == "1"
IS_WINDOWS = platform.system() == "Windows"

bp = Blueprint("main", __name__)
BASE_DIR     = os.path.dirname(os.path.dirname(__file__))
DATA_DIR     = os.path.join(BASE_DIR, "data")
REPORTS_DIR  = os.path.join(BASE_DIR, "reports")
TPL_DATA_PATH = os.path.join(DATA_DIR, "GenoPilot_report_template.docx")
TPL_APP_PATH  = os.path.join(os.path.dirname(__file__), "templates", "report_template.docx")
os.makedirs(REPORTS_DIR, exist_ok=True)

# ============================================================================
# Carga de datos (desde tu Excel ya volcados a JSON)
# ============================================================================
with open(os.path.join(DATA_DIR, "markers.json"), "r", encoding="utf-8") as f:
    MARKERS = json.load(f)
with open(os.path.join(DATA_DIR, "cyp2d6_stars.json"), "r", encoding="utf-8") as f:
    CYP_STARS = json.load(f)
with open(os.path.join(DATA_DIR, "cyp2d6_pheno.json"), "r", encoding="utf-8") as f:
    CYP_PHENO = json.load(f)
with open(os.path.join(DATA_DIR, "recs.json"), "r", encoding="utf-8") as f:
    RECS = json.load(f)

# ============================================================================
# Conjuntos de alelos * para selectores (DPYD / UGT1A1) derivados de MARKERS
# ============================================================================
def _collect_stars(gene: str, extra: list[str] | None = None) -> list[str]:
    """
    Recolecta valores 'star' del JSON de marcadores para poblar selects de diplotipo.
    Conserva '*1' y añade extras si procede; devuelve en orden humano.
    """
    stars = {"*1"}
    for m in MARKERS.get(gene, []):
        s = str(m.get("star", "")).strip()
        if not s or s == "-":
            continue
        s = s.split()[0]  # primer token
        stars.add(s)
    if extra:
        stars.update(extra)
    # orden humano (*1, *2A, *13, HapB3, …)
    return sorted(stars, key=lambda x: (x.replace("*", "0"), x))

DPYD_STARS   = _collect_stars("DPYD")
UGT1A1_STARS = _collect_stars("UGT1A1", extra=["*28"])  # aseguramos *28 visible

# ============================================================================
# Helpers de recomendación corta / limpieza de texto
# ============================================================================
def rtext_short(gene: str, pheno: str, long_text: str) -> str:
    """
    Devuelve texto breve de recomendación para el informe.
    Si no hay match en recs.json, aplica un fallback informativo y seguro.
    """
    SHORT_REC = {
        ("DPYD", "metabolizador normal"):         "Dosis estándar según ficha técnica.",
        ("DPYD", "metabolizador intermedio"):     "Reducir dosis inicial; titular y monitorizar.",
        ("DPYD", "metabolizador lento"):          "Evitar fluoropirimidinas; valorar alternativas.",
        ("UGT1A1", "metabolizador normal"):       "Dosis estándar según ficha técnica.",
        ("UGT1A1", "metabolizador intermedio"):   "Considerar reducción dosis; monitorizar neutropenia.",
        ("UGT1A1", "metabolizador lento"):        "Reducir dosis inicial (30–50%); monitorización estrecha.",
        ("CYP2D6", "metabolizador normal"):       "Dosis estándar.",
        ("CYP2D6", "metabolizador intermedio"):   "Considerar terapia hormonal alternativa.",
        ("CYP2D6", "metabolizador pobre"):        "Evitar tamoxifeno; alternativa terapéutica.",
        ("CYP2D6", "metabolizador ultrarrápido"): "Valorar alternativas según contexto.",
    }
    key = (gene, pheno.lower())
    if key in SHORT_REC:
        return SHORT_REC[key]

    # Limpieza básica cuando hay texto largo en recs.json
    t = re.sub(r"\[.*?\]", "", long_text or "")
    t = re.sub(r"https?://\S+", "", t)
    t = re.sub(r"\s+", " ", t).strip()

    # Fallback robusto (cuando no hay entrada en recs.json)
    if not t:
        return ("No hay recomendación específica en las guías cargadas para este "
                "fenotipo/diplotipo. Seguir ficha técnica y monitorizar estrechamente.")
    return t

# ============================================================================
# Reglas de fenotipo – DPYD
# ============================================================================
def dpyd_from_markers(values: dict):
    """
    Determina diplotipo y fenotipo DPYD a partir de genotipos por marcador.
    Regla docente: 1 variante no-*1 → intermedio; 2 variantes → lento.
    """
    stars_list, polym = [], []
    for m in MARKERS.get("DPYD", []):
        col, ref, var, star = m["column"], m["ref"], m["var"], m["star"]
        geno = values.get(col, "-/-")
        polym.append(f"{col} ({m['rsid']}): {geno}")

        alt = re.split(r"[\/,\s]+", var)[0] if var else ""
        if alt and alt != "-":
            if geno in (f"{ref}/{alt}", f"{alt}/{ref}"):
                stars_list.append(star)             # heterocigosis → 1 evento
            elif geno == f"{alt}/{alt}":
                stars_list.extend([star, star])     # homocigosis → 2 eventos

    if not stars_list:
        dipl, pheno = "*1/*1", "Metabolizador normal"
    elif len(stars_list) == 1:
        dipl, pheno = f"*1/{stars_list[0]}", "Metabolizador intermedio"
    else:
        dipl, pheno = f"{stars_list[0]}/{stars_list[1]}", "Metabolizador lento"

    rec_row = next((r for r in RECS if r["Gene"] == "DPYD" and r["Phenotype"].lower() in pheno.lower()), None)
    return dipl, pheno, rtext_short("DPYD", pheno, (rec_row or {}).get("RecText")), polym


def dpyd_from_diplotype(a1: str, a2: str):
    """
    Determina DPYD directamente desde diplotipo (simple para docencia).
    """
    non1 = sum(1 for x in (a1, a2) if x != "*1")
    if non1 == 0:
        pheno = "Metabolizador normal"
    elif non1 == 1:
        pheno = "Metabolizador intermedio"
    else:
        pheno = "Metabolizador lento"

    dipl = f"{a1}/{a2}"
    rec_row = next((r for r in RECS if r["Gene"] == "DPYD" and r["Phenotype"].lower() in pheno.lower()), None)
    return dipl, pheno, rtext_short("DPYD", pheno, (rec_row or {}).get("RecText")), [f"DPYD diplotipo manual: {dipl}"]

# ============================================================================
# Reglas de fenotipo – UGT1A1
# ============================================================================
def ugt1a1_from_markers(geno: str):
    """
    Interpreta UGT1A1 desde rs887829 (tag de *28): C/C→*1/*1; C/T→*1/*28; T/T→*28/*28.
    """
    polym = [f"UGT1A1*80 (rs887829): {geno}"]
    if geno == "C/C":
        dipl, pheno = "*1/*1", "Metabolizador normal"
    elif geno in ("C/T", "T/C"):
        dipl, pheno = "*1/*28", "Metabolizador intermedio"
    elif geno == "T/T":
        dipl, pheno = "*28/*28", "Metabolizador lento"
    else:
        dipl, pheno = "-/-", "Indeterminado"

    rec_row = next((r for r in RECS if r["Gene"] == "UGT1A1" and r["Phenotype"].lower() in pheno.lower()), None)
    return dipl, pheno, rtext_short("UGT1A1", pheno, (rec_row or {}).get("RecText")), polym


def ugt1a1_from_diplotype(a1: str, a2: str):
    """
    Determina fenotipo UGT1A1 a partir de diplotipo.
    Alelos ↓ función: *28, *80, *6, *37 (incluye *80 como tag en fuerte LD con *28).
    0 ↓ → normal; 1 ↓ → intermedio; ≥2 ↓ → lento.
    """
    decreased = {"*28", "*80", "*6", "*37"}
    pair = "/".join(sorted([a1, a2]))
    count = int(a1 in decreased) + int(a2 in decreased)
    if count == 0:
        pheno = "Metabolizador normal"
    elif count == 1:
        pheno = "Metabolizador intermedio"
    else:
        pheno = "Metabolizador lento"

    rec_row = next((r for r in RECS if r["Gene"] == "UGT1A1" and r["Phenotype"].lower() in pheno.lower()), None)
    return pair, pheno, rtext_short("UGT1A1", pheno, (rec_row or {}).get("RecText")), [f"UGT1A1 diplotipo manual: {pair}"]

# ============================================================================
# Reglas de fenotipo – CYP2D6
# ============================================================================
def _cyp_lookup_pheno(dipl: str) -> str:
    """
    Busca fenotipo de CYP2D6 por diplotipo en la tabla docente; si no, heurística por AS.
    """
    row = next((r for r in CYP_PHENO if r["CYP2D6 Diplotype"] == dipl), None)
    if row:
        return row["Coded Diplotype/Phenotype Summary"]

    # Heurística sencilla por Activity Score (docente)
    def ascore(star: str) -> float:
        loss    = {"*3", "*4", "*5", "*6", "*7", "*14", "*15", "*19", "*59"}
        reduced = {"*10", "*17", "*29", "*41", "*56B"}
        if star in loss:    return 0.0
        if star in reduced: return 0.5
        return 1.0

    a1, a2 = dipl.split("/")
    s = ascore(a1) + ascore(a2)
    if s == 0:       return "Poor Metabolizer"
    if s <= 1.0:     return "Intermediate Metabolizer"
    if s <= 2.25:    return "Normal Metabolizer"
    return "Ultrarapid Metabolizer"


def cyp2d6_from_stars(a1: str, a2: str):
    """
    Interpreta CYP2D6 desde diplotipo de estrellas.
    """
    dipl  = f"{a1}/{a2}"
    pheno = _cyp_lookup_pheno(dipl)
    conv  = {
        "Normal":      "Metabolizador normal",
        "Intermediate":"Metabolizador intermedio",
        "Poor":        "Metabolizador pobre",
        "Ultrarapid":  "Metabolizador ultrarrápido",
    }
    for k, v in conv.items():
        if pheno.startswith(k):
            pheno = v
    rec_row = next(
        (r for r in RECS if r["Gene"] == "CYP2D6" and
         ((pheno.split()[1] in r.get("Phenotype", "")) if " " in pheno else (pheno in r.get("Phenotype", "")))),
        None
    )
    return dipl, pheno, rtext_short("CYP2D6", pheno, (rec_row or {}).get("RecText") or ""), [f"CYP2D6 diplotipo manual: {dipl}"]


def cyp2d6_from_markers(values: dict):
    """
    Construye diplotipo CYP2D6 a partir de marcadores individuales (aprox. docente).
    """
    hits, polym = [], []
    for m in MARKERS.get("CYP2D6", []):
        col, ref, var, star = m["column"], m["ref"], m["var"], str(m["star"]).split()[0]
        geno = values.get(col, "-/-")
        polym.append(f"{col} ({m['rsid']}): {geno}")
        alt = re.split(r"[\/,\s]+", var)[0] if var else ""
        if not star or star == "-":
            continue
        if alt and alt != "-":
            if geno in (f"{ref}/{alt}", f"{alt}/{ref}"):
                hits.append(star)
            elif geno == f"{alt}/{alt}":
                hits.extend([star, star])

    dipl  = "*1/*1" if not hits else (f"*1/{hits[0]}" if len(hits) == 1 else f"{hits[0]}/{hits[1]}")
    pheno = _cyp_lookup_pheno(dipl)
    conv  = {
        "Normal":      "Metabolizador normal",
        "Intermediate":"Metabolizador intermedio",
        "Poor":        "Metabolizador pobre",
        "Ultrarapid":  "Metabolizador ultrarrápido",
    }
    for k, v in conv.items():
        if pheno.startswith(k):
            pheno = v
    rec_row = next(
        (r for r in RECS if r["Gene"] == "CYP2D6" and
         ((pheno.split()[1] in r.get("Phenotype", "")) if " " in pheno else (pheno in r.get("Phenotype", "")))),
        None
    )
    return dipl, pheno, rtext_short("CYP2D6", pheno, (rec_row or {}).get("RecText") or ""), polym

# ============================================================================
# Construcción de tabla de resultados como subdocumento (preview y plantilla)
# ============================================================================
def _shade(cell, hex_fill: str):
    """Aplica sombreado a una celda (hex RGB sin '#')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def _set_fixed_layout(table):
    """Fuerza layout 'fixed' de tabla para que Word respete anchos de columna."""
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

def _set_col_widths(table, widths_cm):
    """Aplica anchos por columna (cm) a todas las filas."""
    _set_fixed_layout(table)
    widths = [Cm(w) for w in widths_cm]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

def _set_font_size(cell, pt):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(pt)

def build_result_subdoc(tpl: DocxTemplate, summary: list):
    """
    Construye la tabla de resultados para incrustar en la plantilla (subdoc).
    Nota: Los anchos aquí son orientativos; el ajuste final lo realiza
    'fit_results_table' tras renderizar el documento completo.
    """
    sub = tpl.new_subdoc()
    table = sub.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Gen", "Fenotipo", "Fármaco de interés", "Recomendación terapéutica"]
    hdr = table.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = t
        _shade(hdr[i], "DDDDDD")
        _set_font_size(hdr[i], 10)

    for row in summary:
        r = table.add_row().cells
        r[0].text = f"{row['gen']}\n{row['diplotipo']}"
        r[1].text = row["fenotipo"]
        r[2].text = row["drug"]
        r[3].text = row["rec"]

        fen = row["fenotipo"].lower()
        shade = "E6F4E6" if "normal" in fen else ("F6DEDE" if any(k in fen for k in ("intermedio", "pobre", "ultra")) else "FFFFFF")
        _shade(r[1], shade)
        _shade(r[3], shade)

        for c in (r[0], r[1], r[2]):
            _set_font_size(c, 9.5)
        _set_font_size(r[3], 9)

    # Anchos base contenidos (se sobreescriben con el ajuste post-render)
    _set_col_widths(table, [3.0, 4.0, 4.0, 6.5])  # ≈17.5 cm total
    return sub

# ============================================================================
# Ajuste post-render: evita "tabla que desborda" en el DOCX final
# ============================================================================
def _norm(s: str) -> str:
    """Normaliza texto (minúsculas + sin acentos) para comparaciones robustas."""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return s.lower().strip()

def _soft_breaks(s: str) -> str:
    """
    Inserta 'zero-width space' tras separadores comunes para permitir el corte
    de líneas largas (URLs, DOIs, fármacos compuestos, etc.).
    """
    if not s:
        return s
    breakers = "/-_.,;:"
    out = []
    for ch in s:
        out.append(ch)
        if ch in breakers:
            out.append("\u200b")
    return "".join(out)

def fit_results_table(doc) -> None:
    """
    Localiza la tabla de RESULTADO y la ajusta al ANCHO ÚTIL DE PÁGINA.
    - Fuerza layout fijo.
    - Asigna anchos proporcionales: 16% | 22% | 22% | 40%.
    - Aplica soft-wrap en columnas 3 y 4 (fármaco, recomendación).
    """
    # 1) Localizar tabla por cabecera
    def is_results_table(t):
        try:
            hdr = [ _norm(c.text) for c in t.rows[0].cells ]
            return ("gen" in hdr[0]
                    and "fenotipo" in hdr[1]
                    and "farmaco" in hdr[2]
                    and "recomendacion" in hdr[3])
        except Exception:
            return False

    table = next((t for t in doc.tables if is_results_table(t)), None)
    if not table:
        return  # nada que ajustar

    # 2) Ancho útil de página en cm
    sec = doc.sections[0]
    usable_cm = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    # Margen de seguridad por redondeos
    usable_cm *= 0.98

    # 3) Proporciones de columnas (ajustables si lo deseas)
    fracs = (0.16, 0.22, 0.22, 0.40)
    widths_cm = [round(usable_cm * f, 2) for f in fracs]
    # Asegurar que la suma no supera el usable
    overflow = sum(widths_cm) - usable_cm
    if overflow > 0:
        widths_cm[-1] = round(widths_cm[-1] - overflow, 2)

    # 4) Aplicar layout y anchos
    _set_col_widths(table, widths_cm)

    # 5) Soft-wrap para textos “duros” (URLs/DOIs) en fármaco y recomendación
    for r in table.rows[1:]:
        for idx in (2, 3):
            for p in r.cells[idx].paragraphs:
                p.text = _soft_breaks(p.text)

# ============================================================================
# Rutas
# ============================================================================
@bp.route("/", methods=["GET", "POST"])
def index():
    # Inputs por marcador (para selectores)
    dpyd_inputs = {m["column"]: m["options"] for m in MARKERS.get("DPYD", [])}
    ugt_inputs  = {m["column"]: m["options"] for m in MARKERS.get("UGT1A1", [])}
    cyp_inputs  = [(m["column"], m["options"]) for m in MARKERS.get("CYP2D6", [])]

    if request.method == "POST":
        # --------------------------- Datos paciente ---------------------------
        patient = {
            "nombre":     request.form.get("nombre", ""),
            "apellidos":  request.form.get("apellidos", ""),
            "full_name":  (request.form.get("nombre", "") + " " + request.form.get("apellidos", "")).strip(),
            "historia":   request.form.get("historia", ""),
            "sexo":       request.form.get("sexo", "-"),
            "fecha_nac":  request.form.get("fecha_nac", "")
        }
        clinical = {
            "enf_actual":  request.form.get("enf_actual", ""),
            "otras_pat":   request.form.get("otras_pat", ""),
            "tratamiento": request.form.get("tto", ""),
        }

        # ------------------------------ DPYD ---------------------------------
        dpyd_mode = request.form.get("dpyd_mode", "markers")
        if dpyd_mode == "diplotype":
            dpyd_a1 = request.form.get("dpyd_a1", "*1")
            dpyd_a2 = request.form.get("dpyd_a2", "*1")
            dpyd_dipl, dpyd_pheno, dpyd_rec, dpyd_poly = dpyd_from_diplotype(dpyd_a1, dpyd_a2)
            dpyd_vals = {}
        else:
            dpyd_vals = {k: request.form.get(k, "-/-") for k in dpyd_inputs.keys()}
            dpyd_dipl, dpyd_pheno, dpyd_rec, dpyd_poly = dpyd_from_markers(dpyd_vals)
            dpyd_a1 = dpyd_a2 = ""

        # ----------------------------- UGT1A1 --------------------------------
        ugt_mode = request.form.get("ugt_mode", "markers")
        if ugt_mode == "diplotype":
            ugt_a1 = request.form.get("ugt_a1", "*1")
            ugt_a2 = request.form.get("ugt_a2", "*1")
            ugt_dipl, ugt_pheno, ugt_rec, ugt_poly = ugt1a1_from_diplotype(ugt_a1, ugt_a2)
            ugt_geno = "-/-"
        else:
            ugt_col  = next(iter(ugt_inputs.keys()))
            ugt_geno = request.form.get(ugt_col, "-/-")
            ugt_dipl, ugt_pheno, ugt_rec, ugt_poly = ugt1a1_from_markers(ugt_geno)
            ugt_a1 = ugt_a2 = ""

        # ----------------------------- CYP2D6 --------------------------------
        cyp_mode = request.form.get("cyp_mode", "diplotype")
        if cyp_mode == "markers":
            cyp_vals = {m[0]: request.form.get(m[0], "-/-") for m in cyp_inputs}
            cyp_dipl, cyp_pheno, cyp_rec, cyp_poly = cyp2d6_from_markers(cyp_vals)
            cyp_a1 = cyp_a2 = ""
        else:
            cyp_a1 = request.form.get("cyp_a1", "*1")
            cyp_a2 = request.form.get("cyp_a2", "*1")
            cyp_dipl, cyp_pheno, cyp_rec, cyp_poly = cyp2d6_from_stars(cyp_a1, cyp_a2)
            cyp_vals = {}

        preview = [
            {"gen": "DPYD",   "dipl": dpyd_dipl, "pheno": dpyd_pheno, "drug": "Fluorouracilo, capecitabina, tegafur", "rec": dpyd_rec},
            {"gen": "CYP2D6", "dipl": cyp_dipl,  "pheno": cyp_pheno,  "drug": "Tamoxifeno",                            "rec": cyp_rec},
            {"gen": "UGT1A1", "dipl": ugt_dipl,  "pheno": ugt_pheno,  "drug": "Irinotecán",                            "rec": ugt_rec},
        ]

        return render_template(
            "preview.html",
            patient=patient, clinical=clinical, preview=preview,

            # DPYD
            dpyd_mode=dpyd_mode, dpyd_inputs=dpyd_inputs, dpyd_vals=dpyd_vals,
            dpyd_a1=dpyd_a1, dpyd_a2=dpyd_a2, dpyd_stars=DPYD_STARS,

            # UGT1A1
            ugt_mode=ugt_mode, ugt_inputs=ugt_inputs, ugt_geno=ugt_geno,
            ugt_a1=ugt_a1, ugt_a2=ugt_a2, ugt_stars=UGT1A1_STARS,

            # CYP2D6
            cyp_mode=cyp_mode, cyp_inputs=cyp_inputs, cyp_vals=cyp_vals,
            cyp_a1=cyp_a1, cyp_a2=cyp_a2, cyp_stars=CYP_STARS,
        )

    # GET: render del formulario
    return render_template(
        "index.html",
        # DPYD
        dpyd_inputs={m["column"]: m["options"] for m in MARKERS.get("DPYD", [])},
        dpyd_stars=DPYD_STARS,
        # UGT1A1
        ugt_inputs={m["column"]: m["options"] for m in MARKERS.get("UGT1A1", [])},
        ugt_stars=UGT1A1_STARS,
        # CYP2D6
        cyp_inputs=[(m["column"], m["options"]) for m in MARKERS.get("CYP2D6", [])],
        cyp_stars=CYP_STARS,
    )

@bp.route("/generate", methods=["POST"])
def generate():
    """
    Genera informe desde el POST del preview.
    - Renderiza plantilla DOCX con subdocumento de resultados.
    - Ajusta la tabla al ancho útil de página (fix overflow).
    - Intenta convertir a PDF con docx2pdf (Word en Windows). Si falla, entrega DOCX.
    """
    # --------------------------- Paciente / Clínica ---------------------------
    patient = {
        "nombre":     request.form.get("nombre", ""),
        "apellidos":  request.form.get("apellidos", ""),
        "full_name":  (request.form.get("nombre", "") + " " + request.form.get("apellidos", "")).strip(),
        "historia":   request.form.get("historia", ""),
        "sexo":       request.form.get("sexo", "-"),
        "fecha_nac":  request.form.get("fecha_nac", "")
    }
    clinical = {
        "enf_actual":  request.form.get("enf_actual", ""),
        "otras_pat":   request.form.get("otras_pat", ""),
        "tratamiento": request.form.get("tto", ""),
    }

    # ------------------------------ DPYD -------------------------------------
    dpyd_inputs = {m["column"]: m["options"] for m in MARKERS.get("DPYD", [])}
    dpyd_mode = request.form.get("dpyd_mode", "markers")
    if dpyd_mode == "diplotype":
        dpyd_a1 = request.form.get("dpyd_a1", "*1")
        dpyd_a2 = request.form.get("dpyd_a2", "*1")
        dpyd_dipl, dpyd_pheno, dpyd_rec, dpyd_poly = dpyd_from_diplotype(dpyd_a1, dpyd_a2)
    else:
        dpyd_vals = {k: request.form.get(k, "-/-") for k in dpyd_inputs.keys()}
        dpyd_dipl, dpyd_pheno, dpyd_rec, dpyd_poly = dpyd_from_markers(dpyd_vals)

    # ------------------------------ UGT1A1 -----------------------------------
    ugt_inputs = {m["column"]: m["options"] for m in MARKERS.get("UGT1A1", [])}
    ugt_mode = request.form.get("ugt_mode", "markers")
    if ugt_mode == "diplotype":
        ugt_a1 = request.form.get("ugt_a1", "*1")
        ugt_a2 = request.form.get("ugt_a2", "*1")
        ugt_dipl, ugt_pheno, ugt_rec, ugt_poly = ugt1a1_from_diplotype(ugt_a1, ugt_a2)
    else:
        ugt_col  = next(iter(ugt_inputs.keys()))
        ugt_geno = request.form.get(ugt_col, "-/-")
        ugt_dipl, ugt_pheno, ugt_rec, ugt_poly = ugt1a1_from_markers(ugt_geno)

    # ------------------------------ CYP2D6 -----------------------------------
    cyp_inputs = [(m["column"], m["options"]) for m in MARKERS.get("CYP2D6", [])]
    cyp_mode = request.form.get("cyp_mode", "diplotype")
    if cyp_mode == "markers":
        cyp_vals = {m[0]: request.form.get(m[0], "-/-") for m in cyp_inputs}
        cyp_dipl, cyp_pheno, cyp_rec, cyp_poly = cyp2d6_from_markers(cyp_vals)
    else:
        cyp_a1 = request.form.get("cyp_a1", "*1")
        cyp_a2 = request.form.get("cyp_a2", "*1")
        cyp_dipl, cyp_pheno, cyp_rec, cyp_poly = cyp2d6_from_stars(cyp_a1, cyp_a2)

    # --------------------------- Contexto informe ----------------------------
    summary = [
        {"gen": "DPYD",   "diplotipo": dpyd_dipl, "fenotipo": dpyd_pheno, "drug": "Fluorouracilo, capecitabina, tegafur", "rec": dpyd_rec},
        {"gen": "CYP2D6", "diplotipo": cyp_dipl,  "fenotipo": cyp_pheno,  "drug": "Tamoxifeno",                            "rec": cyp_rec},
        {"gen": "UGT1A1", "diplotipo": ugt_dipl,  "fenotipo": ugt_pheno,  "drug": "Irinotecán",                            "rec": ugt_rec},
    ]

    tpl_path = TPL_DATA_PATH if os.path.exists(TPL_DATA_PATH) else TPL_APP_PATH
    tpl = DocxTemplate(tpl_path)
    tabla_subdoc = build_result_subdoc(tpl, summary)

    now = datetime.datetime.now()
    context = {
        "patient": patient,
        "clinical": clinical,
        "TABLA_RESULTADO": tabla_subdoc,
        "polymorphisms": "; ".join(dpyd_poly + cyp_poly + ugt_poly),
        "sources": {"cpic_url": "https://cpicpgx.org/guidelines/", "dpwgd_doi": "DOI:10.1038/s41431-022-01243-2"},
        "meta": {
            "sample_code": patient.get("historia") or "—",
            "request_date": now.strftime("%d/%m/%Y"),
            "report_date": now.strftime("%d/%m/%Y"),
            "generated_at": now.strftime("%Y-%m-%d %H:%M"),
            "last_update": "septiembre 2025",
            "version": "0.5.0",
        },
    }

    # ------------------------- Render + ajuste tabla -------------------------
    # 1) Renderizamos a un DOCX temporal
    tmp_docx = os.path.join(tempfile.gettempdir(), f"GenoPilot_tmp_{now.strftime('%H%M%S')}.docx")
    tpl.render(context)
    tpl.save(tmp_docx)

    # 2) Reabrimos y AJUSTAMOS la tabla al ancho útil de página (evita desborde)
    try:
        doc = Document(tmp_docx)
        fit_results_table(doc)   # <-- fix overflow
        doc.save(tmp_docx)
    except Exception:
        # Si por algún motivo python-docx falla, seguimos con el DOCX renderizado
        pass

    # --------------------------- Exportación final ---------------------------
    pdf_name = f"GenoPilot_{patient.get('full_name','Paciente')}_{now.strftime('%Y%m%d_%H%M')}.pdf"
    pdf_path = os.path.join(REPORTS_DIR, pdf_name)

    if IS_WINDOWS and not DISABLE_PDF:
        try:
            from docx2pdf import convert
            convert(tmp_docx, pdf_path)   # Requiere MS Word (Windows)
            return send_file(pdf_path, as_attachment=True, download_name=pdf_name, mimetype="application/pdf")
        except Exception:
            pass  # Fallback a DOCX más abajo

    # Fallback: entregar DOCX
    return send_file(
        tmp_docx,
        as_attachment=True,
        download_name=pdf_name.replace(".pdf", ".docx"),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

